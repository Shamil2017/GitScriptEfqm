import os
import json
from docx import Document

def load_json(json_file):
    """Загрузка JSON файла."""
    with open(json_file, 'r', encoding='utf-8') as file:
        data = json.load(file)
    return data

def generate_word_document(data, output_file):
    """Генерация Word документа из JSON данных."""
    # Создаем новый документ
    doc = Document()

    # Добавляем название кафедры
    doc.add_heading(f"Кафедра: {data['SubdivisionShortName']}", level=1)

    # Перебираем страницы
    for page in data['pages']:
        doc.add_heading(page['page'], level=2)

        # Перебираем вопросы
        for question in page['questions']:
            doc.add_heading(question['question_description'], level=3)

            # Заголовок таблицы
            table = doc.add_table(rows=1, cols=6)
            table.style = 'Table Grid'
            header_cells = table.rows[0].cells
            header_cells[0].text = "Вариант ответа"
            header_cells[1].text = "Проставленный балл"
            header_cells[2].text = "Предложенный балл"
            header_cells[3].text = "Обоснование краткое"
            header_cells[4].text = "Обоснование подробное"
            header_cells[5].text = "Несоответствие"

            # Разбираем результаты
            results = json.loads(question['results'])
            for result in results:
                row_cells = table.add_row().cells
                row_cells[0].text = result["Вариант ответа"]
                row_cells[1].text = str(result["Проставленный балл"])
                row_cells[2].text = str(result["Предложенный балл"])
                row_cells[3].text = result["Обоснование краткое"]
                row_cells[4].text = result["Обоснование подробное"]
                row_cells[5].text = str(result["Несоответствие"])

    # Сохраняем документ
    doc.save(output_file)
    print(f"Документ успешно создан: {output_file}")

def process_json_files(input_folder, output_folder):
    """Обработка всех JSON файлов в папке."""
    # Создаём папку для результатов, если её нет
    if not os.path.exists(output_folder):
        os.makedirs(output_folder)

    # Перебираем все файлы в папке
    for file_name in os.listdir(input_folder):
        if file_name.endswith('.json'):
            input_file = os.path.join(input_folder, file_name)
            output_file = os.path.join(output_folder, file_name.replace('.json', '.docx'))

            # Загружаем данные и генерируем документ
            try:
                data = load_json(input_file)
                generate_word_document(data, output_file)
            except Exception as e:
                print(f"Ошибка при обработке файла {file_name}: {e}")

# Основная программа
def main():
    input_folder = "results"  # Папка с JSON файлами
    output_folder = "results"  # Папка для сохранения DOCX

    # Обрабатываем файлы
    process_json_files(input_folder, output_folder)

if __name__ == "__main__":
    main()
