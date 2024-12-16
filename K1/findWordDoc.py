import os
from docx import Document

def find_files_with_string(input_folder, string_list_file, output_file):
    """
    Ищет строки из файла в содержимом Word-документов в указанной папке
    и записывает имена файлов, содержащих строки, в выходной файл.
    """
    # Читаем строки из файла string_list_file
    with open(string_list_file, "r", encoding="utf-8") as f:
        search_strings = [line.strip() for line in f.readlines()]

    # Получаем список всех файлов .docx в папке
    word_files = [f for f in os.listdir(input_folder) if f.endswith(".docx")]

    # Словарь для хранения результатов
    results = {string: [] for string in search_strings}

    # Перебираем все Word-документы
    for word_file in word_files:
        file_path = os.path.join(input_folder, word_file)
        try:
            doc = Document(file_path)

            # Обходим весь текст (абзацы и таблицы) в документе
            file_content = []
            for paragraph in doc.paragraphs:
                file_content.append(paragraph.text)

            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        file_content.append(cell.text)

            # Проверяем, какие строки присутствуют в содержимом файла
            for search_string in search_strings:
                if any(search_string in content for content in file_content):
                    results[search_string].append(word_file)

        except Exception as e:
            print(f"Ошибка при обработке файла {word_file}: {e}")

    # Записываем результаты в файл
    with open(output_file, "w", encoding="utf-8") as output:
        written_files = set()  # Чтобы избежать дублирования имён файлов
        for files in results.values():
            for file in files:
                if file not in written_files:
                    output.write(file + "\n")
                    written_files.add(file)

    print(f"Результаты поиска записаны в файл: {output_file}")


if __name__ == "__main__":
    # Папка с Word-документами
    input_folder = "Word"

    # Файл со строками для поиска
    string_list_file = "output.txt"

    # Файл для записи результатов
    output_file = "listquestion.txt"

    # Проверяем существование папки
    if not os.path.exists(input_folder):
        print(f"Папка {input_folder} не найдена.")
        exit()

    # Проверяем существование файла со строками
    if not os.path.exists(string_list_file):
        print(f"Файл {string_list_file} не найден.")
        exit()

    # Выполняем поиск строк
    find_files_with_string(input_folder, string_list_file, output_file)
