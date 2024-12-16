import json
from docx import Document

def load_json(json_file):
    """Загрузка JSON файла."""
    with open(json_file, 'r', encoding='utf-8') as file:
        data = json.load(file)
    return data

def generate_word_document(data, output_file):
    """Генерация Word документа из JSON данных."""
    # Создаем новый документ
    doc = Document()

    # Добавляем название кафедры
    doc.add_heading(f"Кафедра: {data['SubdivisionShortName']}", level=1)

    # Перебираем страницы
    for page in data['pages']:
        doc.add_heading(page['page'], level=2)

        # Перебираем вопросы
        for question in page['questions']:
            doc.add_heading(question['question_description'], level=3)

            # Заголовок таблицы
            table = doc.add_table(rows=1, cols=6)
            table.style = 'Table Grid'
            header_cells = table.rows[0].cells
            header_cells[0].text = "Вариант ответа"
            header_cells[1].text = "Проставленный балл"
            header_cells[2].text = "Предложенный балл"
            header_cells[3].text = "Обоснование краткое"
            header_cells[4].text = "Обоснование подробное"
            header_cells[5].text = "Несоответствие"

            # Разбираем результаты
            results = json.loads(question['results'])
            for result in results:
                row_cells = table.add_row().cells
                row_cells[0].text = result["Вариант ответа"]
                row_cells[1].text = str(result["Проставленный балл"])
                row_cells[2].text = str(result["Предложенный балл"])
                row_cells[3].text = result["Обоснование краткое"]
                row_cells[4].text = result["Обоснование подробное"]
                row_cells[5].text = str(result["Несоответствие"])

    # Сохраняем документ
    doc.save(output_file)
    print(f"Документ успешно создан: {output_file}")

# Основная программа
def main():
    # Файлы
    input_json = "out_01_12_24_РТС_7O0512KC7E6SI2.json"
    output_docx = "результат_генерация.docx"

    # Загрузка данных из JSON
    data = load_json(input_json)

    # Генерация Word документа
    generate_word_document(data, output_docx)

if __name__ == "__main__":
    main()
