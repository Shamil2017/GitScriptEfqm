import os
from docx import Document
from docx.shared import Pt
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls


def apply_table_style(table):
    """
    Применяет стиль с границами ко всем ячейкам таблицы.
    """
    for row in table.rows:
        for cell in row.cells:
            tc = cell._element
            tc_pr = tc.get_or_add_tcPr()
            tcBorders = parse_xml(
                r"""
                <w:tcBorders %s>
                    <w:top w:val="single" w:sz="4" w:space="0" w:color="000000"/>
                    <w:left w:val="single" w:sz="4" w:space="0" w:color="000000"/>
                    <w:bottom w:val="single" w:sz="4" w:space="0" w:color="000000"/>
                    <w:right w:val="single" w:sz="4" w:space="0" w:color="000000"/>
                </w:tcBorders>
                """ % nsdecls("w")
            )
            tc_pr.append(tcBorders)


def process_word_files(input_folder, list_file, output_file):
    """
    Обрабатывает Word-документы (.docx) в порядке, указанном в list_file, и объединяет их в один файл.
    """
    merged_doc = Document()

    # Читаем список файлов из listquestion.txt
    with open(list_file, "r", encoding="utf-8") as f:
        ordered_files = [line.strip() for line in f.readlines()]

    # Последовательно обрабатываем файлы
    for file_number, file_name in enumerate(ordered_files, start=1):
        file_path = os.path.join(input_folder, file_name)

        if not os.path.exists(file_path):
            print(f"Файл {file_name} пропущен (не найден).")
            continue

        try:
            # Открываем текущий Word-документ
            doc = Document(file_path)

            # Добавляем номер файла в начало страницы
            header = merged_doc.add_paragraph(f"№ {file_number}: {file_name}")
            header.bold = True
            header.runs[0].font.size = Pt(16)

            # Копируем текстовые абзацы
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():  # Пропускаем пустые строки
                    merged_doc.add_paragraph(paragraph.text)

            # Копируем таблицы
            for table in doc.tables:
                # Создаём таблицу с таким же количеством строк и столбцов
                merged_table = merged_doc.add_table(rows=0, cols=len(table.columns))
                for row in table.rows:
                    new_row = merged_table.add_row()
                    for idx, cell in enumerate(row.cells):
                        new_row.cells[idx].text = cell.text

                # Применяем стиль таблицы
                apply_table_style(merged_table)

            # Добавляем разрыв страницы после каждого документа
            merged_doc.add_page_break()

        except Exception as e:
            print(f"Ошибка при обработке файла {file_name}: {e}")

    # Сохраняем объединённый документ
    merged_doc.save(output_file)
    print(f"Объединённый документ сохранён как {output_file}")


if __name__ == "__main__":
    # Папка с Word-документами
    input_folder = "Word"

    # Файл со списком в порядке обработки
    list_file = "listquestion.txt"

    # Имя выходного файла
    output_file = "result.docx"

    # Проверяем существование папки
    if not os.path.exists(input_folder):
        print(f"Папка {input_folder} не найдена.")
        exit()

    # Проверяем существование файла listquestion.txt
    if not os.path.exists(list_file):
        print(f"Файл {list_file} не найден.")
        exit()

    # Обрабатываем Word-документы в указанном порядке
    process_word_files(input_folder, list_file, output_file)
