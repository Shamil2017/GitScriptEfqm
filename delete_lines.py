from docx import Document

def remove_empty_lines(input_file, output_file):
    """
    Удаляет пустые строки из Word-документа и сохраняет результат.
    """
    # Загружаем документ
    doc = Document(input_file)

    # Удаляем пустые параграфы
    for para in doc.paragraphs:
        if not para.text.strip():  # Проверяем, что текст пустой
            p = para._element
            p.getparent().remove(p)

    # Удаляем пустые строки из таблиц
    for table in doc.tables:
        for row in table.rows:
            if all(not cell.text.strip() for cell in row.cells):  # Если все ячейки в строке пусты
                tbl_row = row._element
                tbl_row.getparent().remove(tbl_row)

    # Сохраняем изменения в новый файл
    doc.save(output_file)
    print(f"Обработанный файл сохранён: {output_file}")

# Пути к файлам
input_docx = "input2.docx"  # Укажите путь к входному файлу
output_docx = "output.docx"  # Укажите путь к выходному файлу

# Запуск функции
remove_empty_lines(input_docx, output_docx)
